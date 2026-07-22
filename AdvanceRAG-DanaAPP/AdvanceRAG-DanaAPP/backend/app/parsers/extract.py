import io
import re
import json as _json

import httpx
from pypdf import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup


def _parse_pdf_date(raw: str | None) -> str | None:
    """PDF metadata dates look like D:20230521120000+00'00' — pull out
    just the YYYY-MM-DD part."""
    if not raw:
        return None
    match = re.search(r"D:(\d{4})(\d{2})(\d{2})", raw)
    if not match:
        return None
    year, month, day = match.groups()
    return f"{year}-{month}-{day}"


def _parse_frontmatter(text: str) -> tuple[dict, str]:
    """Small YAML-frontmatter reader for markdown/text files: reads flat
    `key: value` lines between a leading `---` fence, without pulling in a
    full YAML parser dependency. If there's no frontmatter fence, returns
    an empty metadata dict and the original text untouched."""
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end == -1:
        return {}, text
    block = text[3:end]
    body = text[end + 4:].lstrip("\n")
    meta: dict = {}
    for line in block.splitlines():
        if ":" not in line:
            continue
        key, _, value = line.partition(":")
        meta[key.strip().lower()] = value.strip().strip("\"'")
    return meta, body


def extract_text_and_metadata(filename: str, content: bytes) -> tuple[str, dict]:
    """Returns (text, metadata). `metadata` always has `author`,
    `document_date` (as a YYYY-MM-DD string when known) and `title` keys,
    each None when the format doesn't carry that information or it wasn't
    set on the source file."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    metadata = {"author": None, "document_date": None, "title": None}

    if ext in ("txt", "md"):
        text = content.decode("utf-8", errors="ignore")
        fm, body = _parse_frontmatter(text)
        metadata["author"] = fm.get("author")
        metadata["document_date"] = fm.get("date") or fm.get("document_date")
        metadata["title"] = fm.get("title")
        return body, metadata

    if ext == "pdf":
        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        info = reader.metadata or {}
        metadata["author"] = (info.get("/Author") or "").strip() or None
        metadata["title"] = (info.get("/Title") or "").strip() or None
        metadata["document_date"] = _parse_pdf_date(info.get("/CreationDate"))
        return "\n\n".join(pages), metadata

    if ext == "docx":
        doc = DocxDocument(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        props = doc.core_properties
        metadata["author"] = props.author or None
        metadata["title"] = props.title or None
        metadata["document_date"] = props.created.strftime("%Y-%m-%d") if props.created else None
        return "\n".join(parts), metadata

    if ext in ("html", "htm"):
        soup = BeautifulSoup(content.decode("utf-8", errors="ignore"), "html.parser")
        author_tag = soup.find("meta", attrs={"name": "author"})
        date_tag = soup.find("meta", attrs={"name": "date"}) or soup.find(
            "meta", attrs={"property": "article:published_time"}
        )
        if author_tag:
            metadata["author"] = (author_tag.get("content") or "").strip() or None
        if date_tag:
            metadata["document_date"] = (date_tag.get("content") or "").strip()[:10] or None
        if soup.title and soup.title.get_text().strip():
            metadata["title"] = soup.title.get_text().strip()

        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        return re.sub(r"\n{2,}", "\n\n", text).strip(), metadata

    # fallback: try to decode as plain text, no structured metadata available
    return content.decode("utf-8", errors="ignore"), metadata


def extract_text(filename: str, content: bytes) -> str:
    """Back-compat wrapper for callers that only need the extracted text."""
    text, _ = extract_text_and_metadata(filename, content)
    return text


SUPPORTED_EXTENSIONS = {"pdf", "txt", "md", "html", "htm", "docx"}
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp", "gif"}


_BLOCKED_HOSTS = {"localhost", "127.0.0.1", "0.0.0.0", "::1"}


def _is_safe_url(url: str) -> bool:
    """Minimal SSRF guard: only plain http(s) URLs to a non-loopback host.
    Not a substitute for a real allow-list in a multi-tenant production
    deployment, but blocks the obvious "fetch my own backend" cases."""
    m = re.match(r"^https?://([^/:?#]+)", url, re.IGNORECASE)
    if not m:
        return False
    host = m.group(1).lower()
    if host in _BLOCKED_HOSTS or host.startswith("169.254.") or host.startswith("10.") \
            or host.startswith("192.168.") or re.match(r"^172\.(1[6-9]|2\d|3[0-1])\.", host):
        return False
    return True


def _flatten_json(data, prefix: str = "") -> list[str]:
    """Turns an arbitrary JSON payload (typical of a REST API response) into
    readable `key: value` lines, so it can be chunked/embedded/searched like
    any other text source."""
    lines: list[str] = []
    if isinstance(data, dict):
        for key, value in data.items():
            path = f"{prefix}.{key}" if prefix else str(key)
            if isinstance(value, (dict, list)):
                lines.extend(_flatten_json(value, path))
            else:
                lines.append(f"{path}: {value}")
    elif isinstance(data, list):
        for i, item in enumerate(data):
            path = f"{prefix}[{i}]"
            if isinstance(item, (dict, list)):
                lines.extend(_flatten_json(item, path))
            else:
                lines.append(f"{path}: {item}")
    else:
        lines.append(f"{prefix}: {data}")
    return lines


async def extract_from_url(url: str, timeout: float = 30.0) -> tuple[str, dict, str]:
    """Fetches a web page or API endpoint and returns (text, metadata,
    source_type) — the same shape `extract_text_and_metadata` returns for
    uploaded files, so it can be dropped straight into the same
    chunk/embed/index pipeline. Covers the "websites" and "APIs" ingestion
    sources from the architecture: HTML pages get the same readable-text
    extraction as uploaded .html files; JSON API responses are flattened
    into readable key/value lines; anything else falls back to plain text.
    """
    if not _is_safe_url(url):
        raise ValueError("این آدرس مجاز نیست (فقط http/https و میزبان‌های عمومی پشتیبانی می‌شود).")

    async with httpx.AsyncClient(follow_redirects=True, timeout=timeout) as client:
        resp = await client.get(url, headers={"User-Agent": "dana-rag-bot/1.0"})
        resp.raise_for_status()

    content_type = resp.headers.get("content-type", "").lower()
    metadata = {"author": None, "document_date": None, "title": url}

    if "application/json" in content_type or "+json" in content_type:
        try:
            data = resp.json()
            text = "\n".join(_flatten_json(data))
        except Exception:
            text = resp.text
        return text, metadata, "api"

    if "html" in content_type or resp.text.lstrip().startswith(("<!DOCTYPE", "<html")):
        text, html_metadata = extract_text_and_metadata("page.html", resp.content)
        metadata["author"] = html_metadata.get("author")
        metadata["document_date"] = html_metadata.get("document_date")
        metadata["title"] = html_metadata.get("title") or url
        return text, metadata, "web"

    # Fallback: treat as plain text (covers text/plain, markdown, csv, etc.)
    return resp.text, metadata, "web"
